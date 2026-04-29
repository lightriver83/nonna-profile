#!/usr/bin/env python3
"""Genera Nana_Site_Texts.docx — tutti i testi del sito, divisi per capitoli,
pronti per revisione di gruppo. Ogni testo riporta tra parentesi quadre
l'ID della sezione / elemento, così lo sviluppatore sa dove riportare la modifica nell'HTML."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/nicolagiunchi/Repository/nonna-profile/Nana_Site_Texts.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def H1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x8B, 0x3A, 0x1F)
    return p


def H2(text):
    return doc.add_heading(text, level=2)


def H3(text):
    return doc.add_heading(text, level=3)


def P(text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def LABEL(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run(value)
    return p


def SEP():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— • —")
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def NOTE(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r.font.size = Pt(10)


# ========== COPERTINA ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("NANA")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0x8B, 0x3A, 0x1F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Preserve tradition, one recipe at a time.")
r.italic = True
r.font.size = Pt(16)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run("Testi completi del sito — documento di lavoro per revisione di gruppo")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
NOTE("Come usare questo documento: ogni capitolo corrisponde a una sezione del sito. "
     "Modificate direttamente i testi qui. Il riferimento tra parentesi quadre "
     "[ID sezione / elemento] serve al team di sviluppo per riportare le modifiche "
     "nell'HTML. I capitoli finali (16 e 17) coprono le due pagine secondarie "
     "(catalogo ricette + dettaglio ricetta).")
NOTE("Brand: Nana (4 lettere, niente doppie). Non confondere con 'Nonna Teresa' "
     "e simili, che sono i nomi delle nonne-persone e restano invariati.")
NOTE("Ordine delle sezioni nella home: Hero → Story → Mission → Vision → How → "
     "Recipes preview (3 card) → Community → Team → Why Now → Trust → "
     "By the Numbers (timeline + grafici, in fondo) → CTA → Debrief → Footer.")

doc.add_page_break()

# ========== META / SEO ==========
H1("Capitolo 0 — Meta & SEO (index.html)")
H3("[head / title]")
P("Nana — Preserve tradition, one recipe at a time.")
H3("[head / meta description]")
P("Nana is the community cookbook. Record your grandmother's recipes, share them "
  "with the world, cook each other's Sunday lunch.")
H3("[head / Open Graph title]")
P("Nana — Preserve tradition, one recipe at a time.")
H3("[head / Open Graph description]")
P("A table set in 47 countries. Five founders, five grandmothers, a thousand recipes to share.")

doc.add_page_break()

# ========== NAVIGAZIONE ==========
H1("Capitolo 1 — Navigazione (header)")
H3("[nav / logo]"); P("Nana.")
H3("[nav / link 1]"); P("Story")
H3("[nav / link 2]"); P("Mission")
H3("[nav / link 3]"); P("How")
H3("[nav / link 4]"); P("Recipes")
H3("[nav / link 5]"); P("Community")
H3("[nav / link 6]"); P("Team")
H3("[nav / CTA button]"); P("Join the table")

doc.add_page_break()

# ========== HERO ==========
H1("Capitolo 2 — Hero (sezione di apertura)")

H3("[hero / kicker]")
P("Founded in Bologna · Cooking with the world since 2024")

H3("[hero / titolo principale]")
P("Preserve tradition, one recipe at a time.")
NOTE("Nel sito 'tradition' è in corsivo, e il titolo è spezzato su 4 righe animate: "
     "Preserve / tradition, / one recipe / at a time.")

H3("[hero / meta block]")
LABEL("Company", "Nana S.r.l.")
LABEL("Home of", "The community cookbook")
LABEL("Reach", "5 continents · 47 countries")

H3("[hero / alt text immagine]")
P("Italian grandmother in a cream apron rolling pasta dough at a rustic wooden "
  "table, warm golden light, copper pots in the background.", italic=True)

H3("[hero / scroll indicator]")
P("Scroll")

doc.add_page_break()

# ========== MARQUEE ==========
H1("Capitolo 3 — Marquee (banda scorrevole)")
NOTE("Testi in loop orizzontale sotto l'hero. Separati da pallini •.")
P("• The community cookbook")
P("• From your nonna's hands to everyone's table")
P("• 47 countries · 612 regional cuisines · one Sunday lunch")
P("• Preserve tradition, one recipe at a time")

doc.add_page_break()

# ========== STORY ==========
H1("Capitolo 4 — Story (Our Story)")
NOTE("La timeline è stata spostata in fondo alla pagina (Capitolo 11 — By the Numbers). "
     "Questa sezione contiene solo il racconto origin story.")

H3("[story / eyebrow]")
P("01 · Our Story")

H3("[story / titolo sezione]")
P("A recipe that started it all.")
NOTE("'A recipe' è in corsivo nel sito.")

H3("[story / didascalia foto nonna]")
P("Nonna Maura — the woman whose cappelletti notebook started all of this.")

H3("[story / paragrafo 1 — drop cap]")
P("One Sunday morning in Romagna, five friends pulled up chairs around the same "
  "wooden table in Bologna. Between them, they had brought a yellowed recipe "
  "notebook — tomato stains, handwriting that leaned like a vine — borrowed "
  "from a grandmother's kitchen drawer. They cooked from it together. None of "
  "them went home on time.")

H3("[story / paragrafo 2]")
P("Those five friends — Nicola Giunchi, Vittoria Marcacci, Chiara Ottolini, "
  "Nicola Giovannelli, and Alessia Diodato — became five co-founders that same "
  "evening. Each one brought their own family kitchen: Nonna Maura from the Marche, "
  "Nonna Paola from Emilia Romagna, Nonna Rosa from Basilicata, Nonna Irlanda "
  "from Lazio, and Nonno Nello from Campania. Five very different kitchens, one "
  "identical feeling — being somewhere that smelled like belonging.")

H3("[story / paragrafo 3]")
P("So the five of them built what they wished they'd always had: a simple, "
  "beautiful way to record a grandmother cooking, share her recipes with anyone "
  "curious enough to cook them, and let her table travel.")

H3("[story / callout finale]")
P("— Nana was born on a long wooden table in Bologna, on 17 February 2024.")

doc.add_page_break()

# ========== MISSION ==========
H1("Capitolo 5 — Mission")
H3("[mission / eyebrow]"); P("02 · Mission")
H3("[mission / titolo grande]")
P("To feel at home with your grandmother, anywhere in the world.")
NOTE("'at home' in corsivo.")
H3("[mission / sottotitolo]")
P("One recipe, one table, one memory at a time.")

doc.add_page_break()

# ========== VISION ==========
H1("Capitolo 6 — Vision")
H3("[vision / eyebrow]"); P("03 · Vision")
H3("[vision / titolo grande]")
P("A world where every kitchen holds a grandmother — and every grandchild has a thousand.")
NOTE("'a grandmother' in corsivo.")
H3("[vision / sottotitolo]")
P("By 2035: one million recipes shared across five continents, one Sunday lunch at a time.")

doc.add_page_break()

# ========== HOW IT WORKS ==========
H1("Capitolo 7 — How it works (4 step)")

H3("[how / eyebrow]")
P("04 · How it works")

H3("[how / titolo sezione]")
P("Four steps. Zero complications.")
NOTE("'Zero complications.' in corsivo.")

H3("[how / lede — testo introduttivo]")
P("Grandparents don't need another app. They need to be listened to. So we "
  "designed Nana around the most natural act there is: cooking for someone "
  "you love — then letting them cook it for someone else.")

H3("[how / step 01 — titolo]"); P("Record")
H3("[how / step 01 — descrizione]")
P("A grandchild presses one button. Our guided overlay captures every step in "
  "cinema-grade video, with automatic multilingual subtitles. No tripods, no "
  "jargon, no editing.")

H3("[how / step 02 — titolo]"); P("Upload")
H3("[how / step 02 — descrizione]")
P("The recipe becomes an entry in your private Family Cookbook — video, "
  "handwritten card, voice note, ingredients, the story behind it. Kept like "
  "an heirloom, not a file.")

H3("[how / step 03 — titolo]"); P("Share")
H3("[how / step 03 — descrizione]")
P("With one tap, publish the recipe to the community cookbook. Other grandchildren "
  "from anywhere in the world discover it, save it, cook it, and tell you how their "
  "Sunday went. Your nonna's table just got bigger.")

H3("[how / step 04 — titolo]"); P("Taste")
H3("[how / step 04 — descrizione]")
P("One click ships the artisan ingredient kit to your door — or to your mother's, "
  "your best friend's, your long-distance cousin's. Nonna's exact proportions, "
  "rebuilt in their kitchen.")

doc.add_page_break()

# ========== RECIPES PREVIEW (3 card, dopo How) ==========
H1("Capitolo 8 — The community cookbook (preview · 3 ricette)")
NOTE("Questa sezione arriva subito dopo 'How it works' per mostrare al volo la "
     "customer experience. Solo 3 card cliccabili come anteprima del catalogo completo.")

H3("[recipes / eyebrow]")
P("05 · The community cookbook")

H3("[recipes / titolo sezione]")
P("Taste it before you cook it.")
NOTE("'Taste it' in corsivo.")

H3("[recipes / lede]")
P("Three recipes, three grandmothers, three different Sundays. Click any card "
  "— read the story, watch the steps, ship the basket home and cook it next "
  "Saturday. This is what every recipe in the community feels like.")

SEP()

for nome, by, cat, seats, licked in [
    ("Cappelletti romagnoli in brodo", "Nonna Teresa · Forlì, Romagna", "Primi", "412", "89"),
    ("Tiramisù (the real one)", "Nonna Marilena · Treviso, IT", "Dolci", "904", "211"),
    ("Arancini al ragù", "Nonna Rosalia · Palermo, IT", "Street food", "352", "98"),
]:
    H3(f"[recipe-card / {nome}]")
    LABEL("Nome", nome)
    LABEL("Firma", by)
    LABEL("Categoria badge", cat)
    LABEL("🍷 Seats at the table", seats)
    LABEL("😋 Licked clean", licked)

H3("[recipes-preview / CTA finale]")
P("See all 612 recipes →")

NOTE("Metriche sotto ogni card (sostituto del 'like'): "
     "🍷 Seats at the table (persone che hanno cucinato la ricetta) + "
     "😋 Licked clean (persone che hanno finito il piatto fino all'ultima briciola).")

doc.add_page_break()

# ========== COMMUNITY ==========
H1("Capitolo 9 — The community (Match / Meet / Reward)")

H3("[community / eyebrow]")
P("06 · The community")

H3("[community / titolo sezione]")
P("A table is the oldest social network.")
NOTE("'the oldest' in corsivo.")

H2("Blocco I — Match")
H3("[community / block Match / titolo]"); P("Match")
H3("[community / block Match / testo]")
P("Every tortellino has a twin in Tokyo. Our matching engine finds recipes that "
  "share the same heart across different traditions — gnocchi ↔ kopytka ↔ pierogi "
  "— and shows you what each nonna does differently. Two kitchens, one conversation.")
H3("[community / block Match / match-viz]")
LABEL("Card sinistra", "Cappelletti · Romagna")
LABEL("Linea etichetta", "same heart, different grandmother")
LABEL("Card destra", "Gyōza · Kyoto")

H2("Blocco II — Meet")
H3("[community / block Meet / titolo]"); P("Meet")
H3("[community / block Meet / testo]")
P("Tables that cross borders. Once a month, grandchildren in the same city "
  "gather at one of their homes and cook a recipe from the community cookbook "
  "together. Strangers on a Wednesday. Friends by dessert. The nonnas are very pleased.")

H2("Blocco III — Reward")
H3("[community / block Reward / titolo]"); P("Reward")
H3("[community / block Reward / testo]")
P("Every recipe you share feeds another family. For each recipe you publish, "
  "you earn Sunday Credits — spend them on someone else's ingredient basket. "
  "It's not gamification. It's reciprocity, the way a tray of lasagna used to "
  "travel between neighbours.")
H3("[community / block Reward / elenco reazioni]")
LABEL("🍷 Seats at the table", "cooked by others")
LABEL("😋 Licked clean", "nothing left on the plate")
LABEL("🧑‍🍳 Cooked it", "tried in their own kitchen")
LABEL("✉️ Shared with family", "forwarded to a relative")

doc.add_page_break()

# ========== TEAM ==========
H1("Capitolo 10 — Team (5 founders)")

H3("[team / eyebrow]")
P("07 · Our founders")

H3("[team / titolo sezione]")
P("Five founders. Five family kitchens. One table to share them all.")
NOTE("'One table' in corsivo.")

H3("[team / lede]")
P("From the Marche coast to the Campania hills, we grew up in five very different "
  "kitchens. The only thing we had in common was a grandparent who fed us — and a "
  "recipe we can't stop telling anyone about.")

NOTE("Ruoli founder: CEO Nicola Giunchi · CMO Vittoria · COO Chiara · CFO Nicola "
     "Giovannelli · CIO Alessia. Alessia ha un *nonno* (Nello), non una nonna: "
     "il copy della sua card usa esplicitamente 'Nonno Nello'.")

for nome, ruolo, grandma in [
    ("Nicola Giunchi", "Co-founder · CEO",
     "🍝 Nonna Maura · Marche · cappelletti, vincisgrassi, olive all'ascolana."),
    ("Vittoria Marcacci", "Co-founder · CMO",
     "🍅 Nonna Paola · Emilia Romagna · tortellini, ragù della domenica, piadina."),
    ("Chiara Ottolini", "Co-founder · COO",
     "🌶️ Nonna Rosa · Basilicata · peperoni cruschi, lagane e ceci, raschatell."),
    ("Nicola Giovannelli", "Co-founder · CFO",
     "🍳 Nonna Irlanda · Lazio · carbonara, cacio e pepe, coda alla vaccinara."),
    ("Alessia Diodato", "Co-founder · CIO",
     "🍆 Nonno Nello · Campania · parmigiana, ragù napoletano, sfogliatella."),
]:
    H3(f"[team / member {nome}]")
    LABEL("Nome", nome)
    LABEL("Ruolo", ruolo)
    LABEL("Nonna di riferimento", grandma)

H3("[team / pull quote]")
P("\"The magic of a grandmother's kitchen isn't in the recipe. It's in the love "
  "she folds into it — the patience, the memory, the small spell that turns "
  "dinner into belonging. Nana exists to carry that love, one Sunday at a time, "
  "from her table to yours.\"", italic=True)
P("— The five of us · Bologna, 2024")
NOTE("'love' e 'belonging' in corsivo nella pull-quote.")

doc.add_page_break()

# ========== WHY NOW ==========
H1("Capitolo 11 — Why Now")

H3("[whynow / eyebrow]"); P("08 · Why now")

H3("[whynow / titolo sezione]")
P("The best recipes have always travelled by hand.")
NOTE("'travelled by hand.' in corsivo.")

H3("[whynow / immagine — alt]")
P("Italian grandmother's hands on a recipe notebook beside a rolling pin and "
  "flour, afternoon light through a kitchen window.", italic=True)

H3("[whynow / paragrafo 1]")
P("A grandmother's WhatsApp voice note explaining the dough for six minutes. "
  "A sauce-stained post-it on the fridge. A phone call at 6pm — \"how much salt, "
  "really?\". Families have always found a way. Nana just gave that way a home, "
  "a table, and a passport.")

H3("[whynow / paragrafo 2]")
P("We built Nana because technology finally caught up with tenderness — and "
  "because the warmest thing about a grandmother has always been that she cooks "
  "for people she's never met. Now she can.")

doc.add_page_break()

# ========== TRUST ==========
H1("Capitolo 12 — Trust")

H3("[trust / eyebrow]"); P("09 · Trust")

H3("[trust / logo/partner list]")
P("Slow Food · UNESCO Heritage · Financial Times · Monocle · United Ventures · "
  "La Repubblica · Kinfolk · Gambero Rosso")

H3("[trust / press quote]")
P("\"The warmest Italian startup of the decade — and somehow, the most global.\"",
  italic=True)
P("— Financial Times, 12 Nov 2025")

doc.add_page_break()

# ========== BY THE NUMBERS ==========
H1("Capitolo 13 — By the numbers (timeline + grafici, in fondo)")
NOTE("Sezione spostata dopo 'Trust' e prima della CTA. Contiene la timeline "
     "dell'azienda e i 2 grafici SVG (donut categorie + barchart paesi).")

H3("[numbers-block / eyebrow]")
P("10 · By the numbers")

H3("[numbers-block / titolo sezione]")
P("Two years. A thousand kitchens. One long Sunday lunch.")
NOTE("'A thousand kitchens.' in corsivo.")

H3("[numbers-block / lede]")
P("The timeline of a small idea that grew up at a dinner table — and what the "
  "community cookbook looks like today, counted in recipes and continents.")

H2("Timeline — 6 milestones")
NOTE("Ogni milestone ha un'icona emoji, un'etichetta temporale e un testo. L'ultimo "
     "pallino ('Today') è animato con pulse per mostrare 'siamo qui ora'. "
     "Una linea orizzontale progressiva dal primo al sesto pallino connette tutti i "
     "milestone: l'animazione si attiva all'ingresso nello scroll.")
LABEL("📓 Dec 2022", "A recipe notebook rediscovered in Romagna.")
LABEL("🏡 Jun 2023", "First 12 families record recipes across five Italian regions.")
LABEL("🏛️ Feb 2024", "Nana S.r.l. incorporated in Bologna.")
LABEL("💰 Sep 2024", "€1.8M seed round led by United Ventures.")
LABEL("🌍 Mar 2025", "Launch in 47 countries. 37 000 recipes shared.")
LABEL("🪑 Today", "You're invited to the table.")

SEP()

H2("Grafico A — Donut (categorie di ricette)")
H3("[chart-donut / title]"); P("What's cooking right now")
H3("[chart-donut / centro]")
LABEL("Numero", "37,412")
LABEL("Label", "recipes shared")
H3("[chart-donut / legend]")
LABEL("Primi", "28%")
LABEL("Secondi", "22%")
LABEL("Dolci", "24%")
LABEL("Contorni", "12%")
LABEL("Pani & Lievitati", "9%")
LABEL("Street food", "5%")

H2("Grafico B — Barchart orizzontale (paesi per ricette)")
H3("[chart-bars / title]"); P("The table is set in 47 countries")
H3("[chart-bars / bars]")
LABEL("Italy", "38%")
LABEL("Japan", "11%")
LABEL("Greece", "9%")
LABEL("Mexico", "8%")
LABEL("France", "7%")
LABEL("Morocco", "6%")
LABEL("India", "5%")
LABEL("Türkiye", "4%")
LABEL("Others", "12%")
H3("[chart-bars / note]")
P("Top 8 countries by recipes shared — Q1 2026, rolling 90 days.")

doc.add_page_break()

# ========== CTA ==========
H1("Capitolo 14 — CTA finale")

H3("[cta / titolo]")
P("Pull up a chair.")
NOTE("'Pull up' in corsivo.")

H3("[cta / sottotitolo]")
P("Share your grandmother's hands with the world. Or borrow someone else's for "
  "Sunday lunch.")

H3("[cta / button]"); P("Join the table")

doc.add_page_break()

# ========== DEBRIEF DIDATTICO ==========
H1("Capitolo 15 — Classroom Debrief")

H3("[debrief / eyebrow]"); P("11 · Classroom debrief")

H3("[debrief / titolo sezione]")
P("How this profile follows the rules.")
NOTE("'the rules' in corsivo.")

H3("[debrief / lede]")
P("Below — a guided tour of every rule from the course slides and where we "
  "applied it. Use this as a checklist when you build your own company profile.")

SEP()

H2("Rule 1 — Must-have information")
P("• Company name → Nana S.r.l. · brand: Nana (hero + nav)")
P("• Physical address → Bologna, Italy (story + footer)")
P("• Established date → 17 February 2024 (timeline)")
P("• Contact info → CTA + footer")
P("• Sector → The community cookbook (hero meta)")

H2("Rule 2 — Mission vs Vision")
P("Mission: \"to feel at home with your grandmother, anywhere in the world.\" "
  "— one sentence, warm, plausible, inspirational.")
P("Vision: \"a world where every kitchen holds a grandmother, and every grandchild "
  "has a thousand.\" — one sentence, ambitious, unique to us.")

H2("Rule 3 — Strong action words")
P("We deliberately avoided weak verbs (make, do, help). We used:")
P("Record · Upload · Share · Taste · Match · Meet · Reward · Travel")

H2("Rule 4 — Impressive nouns")
P("We never call Nana \"a service\" or \"a platform.\" We use identity nouns: "
  "the community cookbook, a table that travels, a passport for Sunday lunch.")

H2("Rule 5 — Comparatives & superlatives")
P("\"The warmest Italian startup of the decade — and somehow, the most global.\" (press)")
P("\"A table is the oldest social network.\" (community section)")

H2("Rule 6 — Storytelling")
P("The profile opens with five friends around a table in Bologna and a yellowed "
  "recipe notebook from Romagna — a concrete, human origin story rooted in one "
  "Sunday morning.")

H2("Rule 7 — Charts & facts")
P("• Timeline → five years, six milestones, one arrow (at the end of the page)")
P("• Donut chart → recipe categories in the community cookbook")
P("• Horizontal bars → top 8 countries sharing recipes")
P("• Numbers embedded in cards: seats at the table, licked clean")

H2("Rule 8 — The 7 C's of Communication")
P("1. Clear — one idea per section, one sentence per idea.")
P("2. Correct — consistent register, native-level editing.")
P("3. Complete — Story / Mission / Vision / How / Recipes / Community / Team / Trust / Numbers / CTA.")
P("4. Concrete — cappelletti, not \"pasta products.\" Bologna, not \"Europe.\"")
P("5. Concise — no paragraph longer than four lines.")
P("6. Consideration — warm tone for a warm audience (families, grandchildren).")
P("7. Courteous — we never mock other traditions; every nonna is welcome.")

H2("Rule 9 — Company culture vibe")
P("Chosen culture adjectives: Hospitality, Reciprocity, Craft, Belonging.")
P("Tone: mature and reassuring, like a Sunday lunch — not funky or tech-bro.")
P("Palette and typography follow: terracotta + cream, editorial serif, no gradients, "
  "no emojis in marketing copy.")

H2("Rule 10 — Public speaking reminders")
P("• 55% body, 38% voice, 7% words — deliver the warmth first.")
P("• Repeat the key idea 3 times: \"preserve tradition, one recipe at a time.\"")
P("• Breathe on the dots. Pause after every mission/vision line.")
P("• K.I.S.S. — never read a bullet. Speak to people, not at slides.")
P("• Three key ideas max: story, mission, community. Everything else is backup.")

doc.add_page_break()

# ========== FOOTER ==========
H1("Capitolo 16 — Footer")

H3("[footer / logo]"); P("Nana.")

H3("[footer / indirizzo]")
P("Via San Vitale 47")
P("40125 Bologna, Italy")

H3("[footer / contatti]")
P("hello@nana.kitchen")
P("+39 051 000 000")

H3("[footer / crediti]")
LABEL("Label", "Company profile by")
P("[Your group name here]")
P("English Course · Group project · A.Y. 2025/26")

H3("[footer / reading list]")
LABEL("Label", "Reading list")
P("Starbucks · Tesla · Bloomberg · Patagonia · Airbnb")

H3("[footer / bottom bar]")
P("© 2026 Nana S.r.l. — Fictional example for classroom use.")
P("Designed to make you feel at home.")

doc.add_page_break()

# ========== RECIPES CATALOG PAGE ==========
H1("Capitolo 17 — Pagina catalogo (recipes.html)")
NOTE("Seconda pagina del sito, raggiungibile dai link 'Recipes', 'Join the table', "
     "e dalle card dell'anteprima in home. Mostra l'intero catalogo con filtri per categoria.")

H3("[recipes-page / nav CTA (diverso dall'home)]")
P("← Back to company")

H2("Hero catalogo")
H3("[catalog-hero / titolo]")
P("The community cookbook.")
NOTE("'cookbook.' in corsivo.")
H3("[catalog-hero / sottotitolo]")
P("Every Sunday, somewhere in the world, a nonna is cooking. Browse 37,000+ "
  "recipes shared by grandchildren across 47 countries — click any card to cook "
  "it, save it, or ship the ingredients home.")
H3("[catalog-hero / searchbar placeholder]")
P("Search a recipe, a nonna, a region…")
H3("[catalog-hero / searchbar button]")
P("Search")

H2("Filter chips (categorie)")
NOTE("Chip filter row orizzontale. 'All' è selezionato di default.")
P("All · Primi · Secondi · Contorni · Dolci · Pani & Lievitati · Street food")

H2("Elenco ricette nel catalogo (18 card)")
NOTE("Ogni card ha: foto · categoria badge · nome ricetta · firma nonna · "
     "🍷 seats at the table · 😋 licked clean. Click → recipe.html?slug=<nome>")

for nome, by, cat in [
    ("Cappelletti romagnoli in brodo", "Nonna Teresa · Forlì, Romagna", "Primi"),
    ("Orecchiette con cime di rapa", "Nonna Rosa · Bari, IT", "Primi"),
    ("Risotto alla milanese", "Nonna Adelina · Milano, IT", "Primi"),
    ("Osso buco alla milanese", "Nonna Luigia · Milano, IT", "Secondi"),
    ("Pollo alla cacciatora", "Nonna Carmela · Firenze, IT", "Secondi"),
    ("Saltimbocca alla romana", "Nonna Luisa · Roma, IT", "Secondi"),
    ("Caponata", "Nonna Concetta · Catania, IT", "Contorni"),
    ("Peperonata della domenica", "Nonna Ida · Napoli, IT", "Contorni"),
    ("Tiramisù (the real one)", "Nonna Marilena · Treviso, IT", "Dolci"),
    ("Cantucci alle mandorle", "Nonna Giulia · Siena, IT", "Dolci"),
    ("Cassata siciliana", "Nonna Agata · Palermo, IT", "Dolci"),
    ("Focaccia di Recco", "Nonna Pinuccia · Genova, IT", "Pani"),
    ("Piadina romagnola", "Nonna Mirella · Forlì, IT", "Pani"),
    ("Taralli pugliesi al finocchietto", "Nonna Angela · Bari, IT", "Pani"),
    ("Arancini al ragù", "Nonna Rosalia · Palermo, IT", "Street food"),
    ("Panzerotti pugliesi", "Nonna Marta · Bari, IT", "Street food"),
    ("Trenette al pesto", "Nonna Rita · Genova, IT", "Primi"),
    ("Parmigiana di melanzane", "Nonna Assunta · Napoli, IT", "Secondi"),
]:
    LABEL(nome, f"{by} · {cat}")

doc.add_page_break()

# ========== RECIPE DETAIL PAGE ==========
H1("Capitolo 18 — Pagina dettaglio ricetta (recipe.html)")
NOTE("Mockup di una singola ricetta. Il testo qui sotto è quello della ricetta "
     "di default 'Cappelletti romagnoli in brodo — Nonna Teresa' (ricetta di "
     "magro tradizionale della Romagna). Per le altre ricette, il template è "
     "identico: cambia il contenuto dei campi Nome/Firma/Steps/Story/Ingredients.")

H2("Header ricetta")
H3("[recipe / back link]")
P("Back to recipes")
H3("[recipe / titolo]")
P("Cappelletti romagnoli in brodo")
H3("[recipe / firma]")
P("Nonna Teresa · Forlì, Romagna")
H3("[recipe / meta bar]")
LABEL("Time", "3h 00m")
LABEL("Serves", "6")
LABEL("Difficulty", "Medium")

H2("Media")
NOTE("Nella pagina c'è un'immagine full-width con un play button decorativo "
     "sovrapposto (mockup). Sotto: 4 thumbnail cliccabili che cambiano l'immagine "
     "principale. Quando avremo i video reali delle ricette, basterà sostituire "
     "l'<img> con un <video> HTML5 e il player prenderà il controllo.")
H3("[recipe / thumb 1 alt]"); P("Cappelletti in brodo plated", italic=True)
H3("[recipe / thumb 2 alt]"); P("Fresh pasta dough being rolled", italic=True)
H3("[recipe / thumb 3 alt]"); P("Flour, rolling pin and recipe notebook", italic=True)
H3("[recipe / thumb 4 alt]"); P("Steaming bowl of broth", italic=True)

H2("Reaction bar (sostituto del like)")
H3("[recipe / reactions]")
LABEL("🍷 Seats at the table", "412")
LABEL("😋 Licked clean", "89")
LABEL("🧑‍🍳 Cooked it", "237")
LABEL("✉️ Shared with family", "64")
NOTE("Click incrementa il contatore, persistito in localStorage per utente.")

H2("Story della ricetta (side quote)")
H3("[recipe / story quote]")
P("\"In Romagna i cappelletti sono di magro: solo formaggi, niente carne. Mia "
  "nonna Teresa li faceva il 24 dicembre, con le sorelle e le cugine intorno "
  "al tavolone, e una radio accesa in cucina. Il ripieno — ricotta, raviggiolo, "
  "parmigiano, una grattata di limone — lo chiudeva a mano in minuscoli "
  "cappelletti che stavano due in un cucchiaio. Il brodo partiva alle sette del "
  "mattino, di cappone e manzo. A mezzogiorno di Natale tutta la famiglia sapeva "
  "che c'era solo una domanda possibile: 'quanti?'.\"", italic=True)
P("— Nicola, her grandson")

H2("Steps (6)")
H3("[recipe / step 01]")
P("Make the filling the day before. In a bowl, mix 250g fresh ricotta with 150g "
  "raviggiolo (or stracchino if raviggiolo is hard to find), 150g grated "
  "Parmigiano-Reggiano, one whole egg, a generous grate of fresh nutmeg, the "
  "zest of half an untreated lemon, and a small pinch of salt. It should be "
  "smooth, dense, and smell like Christmas. Cover and rest overnight in the fridge.")

H3("[recipe / step 02]")
P("The dough, Sunday morning. 400g flour 00, 4 eggs, a pinch of salt. Knead by "
  "hand for 15 minutes until elastic and pale. Wrap and rest under a warm bowl "
  "for 30 minutes. No shortcuts.")

H3("[recipe / step 03]")
P("Roll the sfoglia. Thin enough to read the Gazzetta through it. Cut into "
  "3×3 cm squares with a sharp knife or a smooth wheel. Keep the rest covered "
  "— the pasta dries faster than you think.")

H3("[recipe / step 04]")
P("Fill & fold the cappelletto. A pea-sized dot of filling in the centre. Fold "
  "the square into a triangle, pressing the edges firmly. Wrap the long side "
  "around your fingertip, pinch the two tips together, and flip up the top fold: "
  "that's the 'hat'. The first twenty will look shy. The next two hundred won't.")

H3("[recipe / step 05]")
P("The brodo. Half a hen (or a piece of capon, if you can find one), 400g beef "
  "shank, one onion, one carrot, one stalk of celery, one small tomato. Cover "
  "with cold water. Four hours on the lowest possible flame. Skim the foam slowly, "
  "never all at once.")

H3("[recipe / step 06]")
P("Cook and serve. Drop the cappelletti into the gently simmering brodo. Four "
  "minutes, no more. Serve in warmed bowls with a snowfall of Parmigiano. The "
  "only question allowed at the table is 'quanti?' — how many?.")

H2("Ingredienti (sidebar)")
H3("[recipe / ingredients title]")
P("Ingredients · serves 6")
H3("[recipe / ingredients list]")
LABEL("Flour 00", "400 g")
LABEL("Eggs (whole, large)", "4")
LABEL("Fresh ricotta", "250 g")
LABEL("Raviggiolo (or stracchino)", "150 g")
LABEL("Parmigiano-Reggiano 24m", "150 g + extra")
LABEL("Lemon zest (untreated)", "½ lemon")
LABEL("Nutmeg, salt", "q.b.")
LABEL("Hen or capon, for brodo", "½ · ≈ 1 kg")
LABEL("Beef shank", "400 g")
LABEL("Onion, carrot, celery, tomato", "1 each")

H2("Basket CTA (acquisto ingredienti)")
H3("[recipe / basket / eyebrow]")
P("Ship the ingredients")
H3("[recipe / basket / titolo]")
P("Nonna Teresa's Sunday basket")
H3("[recipe / basket / prezzo]")
P("€32")
H3("[recipe / basket / tempi]")
P("Order by Thursday · arrives Saturday morning")
H3("[recipe / basket / button]")
P("Add to my table")
NOTE("Dopo il click: '✓ Added to your table' per 2.4s, poi torna al testo originale.")
H3("[recipe / basket / stats bottom]")
LABEL("412", "seats at the table")
LABEL("89", "licked clean")

H2("Commenti della community (3 thread mock)")
H3("[recipe / comment 1]")
LABEL("Autore", "Hiro M.")
LABEL("Quando · Dove", "2 days ago · Kyoto")
P("My grandmother makes gyōza with almost the same folding gesture. I sent her "
  "this video — she wants to invite Nonna Teresa for tea. 🙏")

H3("[recipe / comment 2]")
LABEL("Autore", "Mariana P.")
LABEL("Quando · Dove", "5 days ago · Oaxaca")
P("Cooked it last Sunday with my kids. The brodo made my kitchen feel like my "
  "grandmother's in 1998. The raviggiolo was the detail — thank you for this.")

H3("[recipe / comment 3]")
LABEL("Autore", "Léa B.")
LABEL("Quando · Dove", "1 week ago · Lyon")
P("The basket arrived on Saturday morning, perfect timing. The raviggiolo was "
  "a revelation. One question — the hen, which age? My butcher asked.")

H2("You might also love… (3 card ricette simili)")
H3("[recipe / similar]")
LABEL("Suggested 1", "Tiramisù (the real one) · Nonna Marilena · Treviso, IT")
LABEL("Suggested 2", "Arancini al ragù · Nonna Rosalia · Palermo, IT")
LABEL("Suggested 3", "Piadina romagnola · Nonna Mirella · Forlì, IT")

doc.add_page_break()

# ========== CHIUSURA ==========
NOTE("Fine del documento. Totale capitoli: 19 (da 0 a 18).")
NOTE("File sorgente: /Users/nicolagiunchi/Repository/nonna-profile/")
NOTE("Pagine sito: index.html · recipes.html · recipe.html")

doc.save(OUT)
print(f"OK → {OUT}")
